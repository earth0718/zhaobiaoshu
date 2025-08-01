#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件生成API接口模块

提供招标文件生成相关的REST API接口，包括：
1. 单文档招标文件生成
2. 批量文档处理
3. 生成状态查询
4. 模型配置管理
"""

import os
import uuid
import asyncio
from datetime import datetime
from typing import Optional, Dict, Any, List
from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
import logging
from docx import Document
from docx.shared import Inches

# 导入核心处理模块
from .processor import process_document
from .batch_processor import process_multiple_documents_async
from ..llm_service.model_manager import model_manager
from ..history.history_manager import history_manager
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from config.multi_file_settings import multi_file_config

# 配置日志
logger = logging.getLogger(__name__)

# 创建路由器
router = APIRouter()

# 全局任务状态存储（生产环境建议使用Redis等持久化存储）
task_status = {}

# 请求模型
class TenderGenerationRequest(BaseModel):
    """招标文件生成请求模型"""
    model_provider: Optional[str] = "ollama"
    quality_level: Optional[str] = "standard"
    include_sections: Optional[List[str]] = None
    custom_requirements: Optional[str] = None

class TextTenderGenerationRequest(BaseModel):
    """文本输入招标文件生成请求模型"""
    text_content: str
    model_provider: Optional[str] = "ollama"
    quality_level: Optional[str] = "standard"
    project_name: Optional[str] = "招标项目"
    include_sections: Optional[List[str]] = None
    custom_requirements: Optional[str] = None

class MultipleTenderGenerationRequest(BaseModel):
    """多文件招标文件生成请求模型"""
    model_provider: Optional[str] = "ollama"
    quality_level: Optional[str] = "standard"
    project_name: Optional[str] = "招标项目"
    include_sections: Optional[List[str]] = None
    custom_requirements: Optional[str] = None

class ModelConfigRequest(BaseModel):
    """模型配置请求模型"""
    module_name: str
    provider: str

class TaskStatusResponse(BaseModel):
    """任务状态响应模型"""
    task_id: str
    status: str  # pending, processing, completed, failed
    progress: int  # 0-100
    message: str
    result: Optional[Dict[str, Any]] = None
    created_at: str
    updated_at: str

# 辅助函数
def save_uploaded_file(upload_file: UploadFile, upload_dir: str = "upload") -> str:
    """保存上传的文件并返回文件路径"""
    # 确保上传目录存在
    os.makedirs(upload_dir, exist_ok=True)
    
    # 生成唯一的文件名
    file_extension = os.path.splitext(upload_file.filename)[1]
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    # 保存文件
    with open(file_path, "wb") as buffer:
        content = upload_file.file.read()
        buffer.write(content)
    
    return file_path

def save_multiple_uploaded_files(upload_files: List[UploadFile], upload_dir: str = "upload") -> List[str]:
    """保存多个上传的文件并返回文件路径列表"""
    file_paths = []
    
    for upload_file in upload_files:
        try:
            file_path = save_uploaded_file(upload_file, upload_dir)
            file_paths.append(file_path)
        except Exception as e:
            logger.error(f"保存文件 {upload_file.filename} 失败: {str(e)}")
            # 清理已保存的文件
            for saved_path in file_paths:
                try:
                    os.remove(saved_path)
                except:
                    pass
            raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")
    
    return file_paths

def update_task_status(task_id: str, status: str, progress: int = 0, message: str = "", result: Dict[str, Any] = None):
    """更新任务状态"""
    if task_id in task_status:
        task_status[task_id].update({
            "status": status,
            "progress": progress,
            "message": message,
            "result": result,
            "updated_at": datetime.now().isoformat()
        })

async def process_document_async(task_id: str, file_path: str, config: Dict[str, Any]):
    """异步处理文档生成招标书"""
    start_time = datetime.now()
    original_filename = os.path.basename(file_path)
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
    
    try:
        update_task_status(task_id, "processing", 10, "开始处理文档...")
        
        # 设置模型配置
        if config.get("model_provider"):
            model_manager.set_current_model("tender_generation", config["model_provider"])
        
        update_task_status(task_id, "processing", 30, "正在解析文档内容...")
        
        # 调用核心处理函数
        result = process_document(file_path)
        
        update_task_status(task_id, "processing", 90, "正在生成最终文档...")
        
        # 计算处理时间
        processing_duration = (datetime.now() - start_time).total_seconds()
        
        # 保存历史记录
        try:
            history_manager.add_record(
                task_id=task_id,
                original_filename=original_filename,
                file_size=file_size,
                model_provider=config.get("model_provider", "unknown"),
                quality_level=config.get("quality_level", "standard"),
                tender_content=result,
                status="completed",
                processing_duration=processing_duration
            )
        except Exception as history_error:
            logger.warning(f"保存历史记录失败: {str(history_error)}")
        
        # 清理临时文件
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # 保存生成的招标文件到download目录
        download_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "download")
        word_dir = os.path.join(download_dir, "word")
        markdown_dir = os.path.join(download_dir, "markdown")
        os.makedirs(word_dir, exist_ok=True)
        os.makedirs(markdown_dir, exist_ok=True)
        
        # 生成文件名（使用时间戳确保唯一性）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"tender_{timestamp}_{task_id[:8]}"
        
        # 保存Markdown文件
        md_filename = f"{base_filename}.md"
        md_path = os.path.join(markdown_dir, md_filename)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(result)
        
        # 生成Word文档
        doc = Document()
        doc.add_heading('招标文件', 0)
        
        # 将结果按段落分割并添加到文档
        paragraphs = result.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                if paragraph.startswith('#'):
                    # 处理标题
                    level = len(paragraph) - len(paragraph.lstrip('#'))
                    title_text = paragraph.lstrip('# ').strip()
                    doc.add_heading(title_text, level)
                else:
                    doc.add_paragraph(paragraph.strip())
        
        # 保存Word文档
        doc_filename = f"{base_filename}.docx"
        doc_path = os.path.join(word_dir, doc_filename)
        doc.save(doc_path)
        
        # 完成任务
        update_task_status(task_id, "completed", 100, "招标文件生成完成", {
            "tender_document": result,
            "file_size": len(result),
            "generation_time": datetime.now().isoformat(),
            "processing_duration": processing_duration,
            "files": {
                "word": {
                    "filename": doc_filename,
                    "download_url": f"/api/tender/download/word/{doc_filename}"
                },
                "markdown": {
                    "filename": md_filename,
                    "download_url": f"/api/tender/download/markdown/{md_filename}"
                }
            }
        })
        
    except Exception as e:
        logger.error(f"处理文档时出错: {str(e)}", exc_info=True)
        
        # 计算处理时间（即使失败）
        processing_duration = (datetime.now() - start_time).total_seconds()
        
        # 保存失败记录到历史
        try:
            history_manager.add_record(
                task_id=task_id,
                original_filename=original_filename,
                file_size=file_size,
                model_provider=config.get("model_provider", "unknown"),
                quality_level=config.get("quality_level", "standard"),
                tender_content="",
                status="failed",
                error_message=str(e),
                processing_duration=processing_duration
            )
        except Exception as history_error:
            logger.warning(f"保存失败历史记录失败: {str(history_error)}")
        
        update_task_status(task_id, "failed", 0, f"处理失败: {str(e)}")
        
        # 清理临时文件
        if os.path.exists(file_path):
            os.remove(file_path)

async def process_multiple_documents_async_task(task_id: str, file_paths: List[str], config: Dict[str, Any]):
    """异步处理多个文档的后台任务"""
    try:
        update_task_status(task_id, "processing", 10, "开始处理多个文档...")
        
        # 定义进度回调函数
        def progress_callback(progress: int, message: str):
            if progress == -1:  # 错误情况
                update_task_status(task_id, "failed", 0, message)
            else:
                # 将进度映射到10-90范围
                mapped_progress = 10 + int(progress * 0.8)
                update_task_status(task_id, "processing", mapped_progress, message)
        
        # 调用批量处理函数
        result_content = await process_multiple_documents_async(
            file_paths, config, progress_callback
        )
        
        update_task_status(task_id, "processing", 90, "正在保存生成的招标文件...")
        
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"tender_{timestamp}_{task_id[:8]}"
        
        # 保存为Markdown文件
        markdown_dir = multi_file_config.get_markdown_output_dir()
        os.makedirs(markdown_dir, exist_ok=True)
        markdown_filename = f"{base_filename}.md"
        markdown_path = os.path.join(markdown_dir, markdown_filename)
        
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(result_content)
        
        # 保存为Word文档
        word_dir = multi_file_config.get_word_output_dir()
        os.makedirs(word_dir, exist_ok=True)
        word_filename = f"{base_filename}.docx"
        word_path = os.path.join(word_dir, word_filename)
        
        # 创建Word文档
        doc = Document()
        
        # 解析Markdown内容并添加到Word文档
        lines = result_content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                # 一级标题
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # 二级标题
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # 三级标题
                doc.add_heading(line[4:], level=3)
            elif line == '---':
                # 分隔线，跳过
                continue
            elif line:
                # 普通段落
                doc.add_paragraph(line)
        
        doc.save(word_path)
        
        # 记录成功历史
        try:
            file_names = [os.path.basename(fp) for fp in file_paths]
            history_manager.save_success_record(
                task_id=task_id,
                file_paths=file_paths,
                file_names=file_names,
                markdown_path=markdown_path,
                word_path=word_path,
                config=config,
                content_preview=result_content[:500] + "..." if len(result_content) > 500 else result_content
            )
        except Exception as history_error:
            logger.warning(f"保存成功历史记录失败: {str(history_error)}")
        
        # 更新任务状态为完成
        update_task_status(task_id, "completed", 100, "多文档招标文件生成完成！", {
            "markdown_file": markdown_filename,
            "word_file": word_filename,
            "download_urls": {
                "markdown": f"/api/tender/download/markdown/{markdown_filename}",
                "word": f"/api/tender/download/word/{word_filename}"
            },
            "file_count": len(file_paths),
            "source_files": [os.path.basename(fp) for fp in file_paths]
        })
        
    except Exception as e:
        logger.error(f"处理多个文档失败: {str(e)}", exc_info=True)
        
        # 记录失败历史
        try:
            history_manager.save_failed_record(
                task_id=task_id,
                file_path=",".join(file_paths),
                error_message=str(e),
                config=config
            )
        except Exception as history_error:
            logger.warning(f"保存失败历史记录失败: {str(history_error)}")
        
        update_task_status(task_id, "failed", 0, f"处理失败: {str(e)}")

async def process_text_content_async(task_id: str, text_content: str, config: Dict[str, Any]):
    """异步处理文本内容生成招标书"""
    start_time = datetime.now()
    
    try:
        update_task_status(task_id, "processing", 10, "开始处理文本内容...")
        
        # 设置模型配置
        if config.get("model_provider"):
            model_manager.set_current_model("tender_generation", config["model_provider"])
        
        update_task_status(task_id, "processing", 30, "正在分析文本内容...")
        
        # 导入处理器模块
        from .processor import process_text_content
        
        # 调用文本处理函数
        result = process_text_content(text_content, config)
        
        update_task_status(task_id, "processing", 90, "正在生成最终文档...")
        
        # 计算处理时间
        processing_duration = (datetime.now() - start_time).total_seconds()
        
        # 保存历史记录
        try:
            history_manager.add_record(
                task_id=task_id,
                original_filename=f"文本输入_{config.get('project_name', '招标项目')}",
                file_size=len(text_content),
                model_provider=config.get("model_provider", "unknown"),
                quality_level=config.get("quality_level", "standard"),
                tender_content=result,
                status="completed",
                processing_duration=processing_duration
            )
        except Exception as history_error:
            logger.warning(f"保存历史记录失败: {str(history_error)}")
        
        # 保存生成的招标文件到download目录
        download_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "download")
        word_dir = os.path.join(download_dir, "word")
        markdown_dir = os.path.join(download_dir, "markdown")
        os.makedirs(word_dir, exist_ok=True)
        os.makedirs(markdown_dir, exist_ok=True)
        
        # 生成文件名（使用时间戳确保唯一性）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"tender_{timestamp}_{task_id[:8]}"
        
        # 保存Markdown文件
        md_filename = f"{base_filename}.md"
        md_path = os.path.join(markdown_dir, md_filename)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(result)
        
        # 生成Word文档
        doc = Document()
        doc.add_heading('招标文件', 0)
        
        # 将结果按段落分割并添加到文档
        paragraphs = result.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                if paragraph.startswith('#'):
                    # 处理标题
                    level = len(paragraph) - len(paragraph.lstrip('#'))
                    title_text = paragraph.lstrip('# ').strip()
                    doc.add_heading(title_text, level)
                else:
                    doc.add_paragraph(paragraph.strip())
        
        # 保存Word文档
        doc_filename = f"{base_filename}.docx"
        doc_path = os.path.join(word_dir, doc_filename)
        doc.save(doc_path)
        
        # 完成任务
        update_task_status(task_id, "completed", 100, "招标文件生成完成", {
            "tender_document": result,
            "file_size": len(result),
            "generation_time": datetime.now().isoformat(),
            "processing_duration": processing_duration,
            "files": {
                "word": {
                    "filename": doc_filename,
                    "download_url": f"/api/tender/download/word/{doc_filename}"
                },
                "markdown": {
                    "filename": md_filename,
                    "download_url": f"/api/tender/download/markdown/{md_filename}"
                }
            }
        })
        
    except Exception as e:
        logger.error(f"处理文本内容时出错: {str(e)}", exc_info=True)
        
        # 计算处理时间（即使失败）
        processing_duration = (datetime.now() - start_time).total_seconds()
        
        # 保存失败记录到历史
        try:
            history_manager.add_record(
                task_id=task_id,
                original_filename=f"文本输入_{config.get('project_name', '招标项目')}",
                file_size=len(text_content),
                model_provider=config.get("model_provider", "unknown"),
                quality_level=config.get("quality_level", "standard"),
                tender_content="",
                status="failed",
                error_message=str(e),
                processing_duration=processing_duration
            )
        except Exception as history_error:
            logger.warning(f"保存失败历史记录失败: {str(history_error)}")
        
        update_task_status(task_id, "failed", 0, f"处理失败: {str(e)}")

# API接口
@router.post("/generate", summary="生成招标书", description="上传文档并生成招标书")
async def generate_tender_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(..., description="要处理的文档文件（支持PDF、DOCX格式）"),
    model_provider: Optional[str] = None,
    quality_level: Optional[str] = "standard"
):
    """生成招标书接口"""
    try:
        # 验证文件格式
        if not file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")
        
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in [".pdf", ".docx"]:
            raise HTTPException(status_code=400, detail="不支持的文件格式，仅支持PDF和DOCX")
        
        # 生成任务ID
        task_id = str(uuid.uuid4())
        
        # 保存上传的文件
        file_path = save_uploaded_file(file)
        
        # 初始化任务状态
        task_status[task_id] = {
            "task_id": task_id,
            "status": "pending",
            "progress": 0,
            "message": "任务已创建，等待处理...",
            "result": None,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        
        # 配置参数 - 如果没有指定模型提供商，使用当前配置的模型
        if model_provider is None:
            model_provider = model_manager.get_current_model("tender_generation")
        
        config = {
            "model_provider": model_provider,
            "quality_level": quality_level
        }
        
        # 添加后台任务
        background_tasks.add_task(process_document_async, task_id, file_path, config)
        
        return {
            "task_id": task_id,
            "message": "任务已创建，正在处理中...",
            "status_url": f"/api/tender/status/{task_id}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建招标文件生成任务失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")

@router.post("/generate_multiple", summary="多文件生成招标书", description="上传多个文档并合并生成一份完整招标书")
async def generate_tender_from_multiple_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(..., description="要处理的文档文件列表（支持PDF、DOCX格式）"),
    model_provider: Optional[str] = None,
    quality_level: Optional[str] = "standard",
    project_name: Optional[str] = "招标项目",
    custom_requirements: Optional[str] = None,
    include_sections: Optional[str] = None
):
    """多文件生成招标书接口"""
    try:
        # 验证文件数量
        if not files or len(files) == 0:
            raise HTTPException(status_code=400, detail="至少需要上传一个文件")
        
        max_files = multi_file_config.get_max_files()
        if len(files) > max_files:
            raise HTTPException(status_code=400, detail=f"最多支持同时上传{max_files}个文件")
        
        # 验证文件格式和大小
        max_file_size_mb = multi_file_config.get_max_file_size_mb()
        max_file_size_bytes = max_file_size_mb * 1024 * 1024
        
        for file in files:
            if not file.filename:
                raise HTTPException(status_code=400, detail="文件名不能为空")
            
            # 验证文件格式
            if not multi_file_config.validate_file_format(file.filename):
                supported_formats = ", ".join(multi_file_config.get_supported_formats())
                raise HTTPException(
                    status_code=400, 
                    detail=f"文件 {file.filename} 格式不支持，仅支持以下格式: {supported_formats}"
                )
            
            # 验证文件大小（如果可以获取）
            if hasattr(file, 'size') and file.size and file.size > max_file_size_bytes:
                raise HTTPException(
                    status_code=400,
                    detail=f"文件 {file.filename} 大小超过限制（最大{max_file_size_mb}MB）"
                )
        
        # 生成任务ID
        task_id = str(uuid.uuid4())
        
        # 保存上传的文件
        file_paths = save_multiple_uploaded_files(files)
        
        # 初始化任务状态
        task_status[task_id] = {
            "task_id": task_id,
            "status": "pending",
            "progress": 0,
            "message": f"任务已创建，准备处理 {len(files)} 个文件...",
            "result": None,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        
        # 配置参数
        if model_provider is None:
            model_provider = model_manager.get_current_model("tender_generation")
        
        config = {
            "model_provider": model_provider,
            "quality_level": quality_level,
            "project_name": project_name or multi_file_config.get_default_project_name(),
            "custom_requirements": custom_requirements
        }
        
        # 处理include_sections参数
        if include_sections:
            try:
                # 如果是字符串，尝试按逗号分割
                if isinstance(include_sections, str):
                    config["include_sections"] = [s.strip() for s in include_sections.split(",") if s.strip()]
                else:
                    config["include_sections"] = include_sections
            except Exception:
                # 如果解析失败，忽略该参数
                pass
        
        # 添加后台任务
        background_tasks.add_task(process_multiple_documents_async_task, task_id, file_paths, config)
        
        return {
            "task_id": task_id,
            "message": f"任务已创建，正在处理 {len(files)} 个文件...",
            "file_count": len(files),
            "file_names": [file.filename for file in files],
            "status_url": f"/api/tender/status/{task_id}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建多文件招标文件生成任务失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")

@router.post("/generate_from_text", summary="从文本生成招标书", description="根据用户输入的文本内容生成招标书")
async def generate_tender_from_text(
    background_tasks: BackgroundTasks,
    request: TextTenderGenerationRequest
):
    """从文本生成招标书接口"""
    try:
        # 验证文本内容
        if not request.text_content or len(request.text_content.strip()) < 10:
            raise HTTPException(status_code=400, detail="文本内容不能为空且至少需要10个字符")
        
        # 生成任务ID
        task_id = str(uuid.uuid4())
        
        # 初始化任务状态
        task_status[task_id] = {
            "task_id": task_id,
            "status": "pending",
            "progress": 0,
            "message": "任务已创建，等待处理...",
            "result": None,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        
        # 配置参数 - 如果没有指定模型提供商，使用当前配置的模型
        model_provider = request.model_provider
        if model_provider is None:
            model_provider = model_manager.get_current_model("tender_generation")
        
        config = {
            "model_provider": model_provider,
            "quality_level": request.quality_level,
            "project_name": request.project_name,
            "include_sections": request.include_sections,
            "custom_requirements": request.custom_requirements
        }
        
        # 添加后台任务
        background_tasks.add_task(process_text_content_async, task_id, request.text_content, config)
        
        return {
            "task_id": task_id,
            "message": "文本招标文件生成任务已创建，正在处理中...",
            "status_url": f"/api/tender/status/{task_id}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建文本招标文件生成任务失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")

@router.get("/status/{task_id}", response_model=TaskStatusResponse, summary="查询任务状态")
async def get_task_status(task_id: str):
    """查询任务状态"""
    if task_id not in task_status:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    return task_status[task_id]

@router.get("/models", summary="获取可用模型列表")
async def get_available_models():
    """获取可用的模型列表"""
    try:
        return {
            "current_models": {
                "tender_generation": model_manager.get_current_model("tender_generation")
            },
            "available_providers": model_manager.get_available_providers(),
            "model_status": {
                provider: model_manager.check_model_availability(provider)
                for provider in model_manager.get_available_providers()
            }
        }
    except Exception as e:
        logger.error(f"获取模型列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取模型列表失败: {str(e)}")

@router.post("/models/switch", summary="切换模型")
async def switch_model(request: ModelConfigRequest):
    """切换当前使用的模型"""
    try:
        model_manager.set_current_model(request.module_name, request.provider)
        return {
            "message": f"已切换 {request.module_name} 模块的模型到 {request.provider}",
            "current_model": model_manager.get_current_model(request.module_name)
        }
    except Exception as e:
        logger.error(f"切换模型失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"切换模型失败: {str(e)}")

@router.get("/health", summary="健康检查")
async def health_check():
    """招标生成服务健康检查"""
    try:
        # 检查模型管理器状态
        current_model = model_manager.get_current_model("tender_generation")
        model_available = model_manager.check_model_availability(current_model)
        
        return {
            "status": "healthy" if model_available else "degraded",
            "timestamp": datetime.now().isoformat(),
            "services": {
                "model_manager": "running",
                "current_model": current_model,
                "model_available": model_available
            },
            "active_tasks": len([t for t in task_status.values() if t["status"] == "processing"])
        }
    except Exception as e:
        logger.error(f"健康检查失败: {str(e)}")
        return {
            "status": "unhealthy",
            "timestamp": datetime.now().isoformat(),
            "error": str(e)
        }

@router.delete("/tasks/{task_id}", summary="删除任务")
async def delete_task(task_id: str):
    """删除指定的任务"""
    if task_id not in task_status:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    del task_status[task_id]
    return {"message": f"任务 {task_id} 已删除"}

@router.get("/tasks", summary="获取所有任务列表")
async def get_all_tasks():
    """获取所有任务的状态列表"""
    return {
        "total_tasks": len(task_status),
        "tasks": list(task_status.values())
    }

@router.get("/download/{file_type}/{filename}", summary="下载生成的招标文件")
async def download_tender_file(file_type: str, filename: str):
    """下载生成的招标文件
    
    Args:
        file_type: 文件类型，支持 'word' 或 'markdown'
        filename: 文件名
    """
    # 验证文件类型
    if file_type not in ["word", "markdown"]:
        raise HTTPException(status_code=400, detail="不支持的文件类型，仅支持 word 或 markdown")
    
    # 构建文件路径
    download_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "download")
    file_dir = os.path.join(download_dir, file_type)
    file_path = os.path.join(file_dir, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 验证文件名格式（安全检查）
    if not filename.startswith("tender_"):
        raise HTTPException(status_code=400, detail="无效的文件名")
    
    # 根据文件类型设置媒体类型
    if file_type == "word":
        if not filename.endswith(".docx"):
            raise HTTPException(status_code=400, detail="Word文件必须以.docx结尾")
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:  # markdown
        if not filename.endswith(".md"):
            raise HTTPException(status_code=400, detail="Markdown文件必须以.md结尾")
        media_type = "text/markdown"
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=media_type
    )