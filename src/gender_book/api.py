from fastapi import APIRouter, HTTPException, UploadFile, File, BackgroundTasks, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Dict, List, Any, Optional
import json
import logging
import uuid
import os
import asyncio
from datetime import datetime
import sys
from docx import Document
from docx.shared import Inches

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from .tender_generator import BidProposalGenerator
from .section_manager import SectionManager
# 导入现有的LLM服务
from src.llm_service import LLMService

router = APIRouter(tags=["gender_book"])
logger = logging.getLogger(__name__)

# 任务状态存储
task_status = {}

class BidProposalGenerationRequest(BaseModel):
    """投标书生成请求模型"""
    tender_document_json: Dict[str, Any]
    model_name: Optional[str] = None
    batch_size: Optional[int] = None
    generate_outline_only: Optional[bool] = False

class BidProposalGenerationRequestWithAttachments(BaseModel):
    """带附件的投标书生成请求模型"""
    tender_document_json: str  # 改为字符串，因为要通过Form传递
    model_name: Optional[str] = None
    batch_size: Optional[int] = None
    generate_outline_only: Optional[bool] = False

class BidProposalGenerationResponse(BaseModel):
    """投标书生成响应模型"""
    success: bool
    task_id: Optional[str] = None
    message: str = ""
    error: str = ""
    data: Dict[str, Any] = {}

class TaskStatusResponse(BaseModel):
    """任务状态响应模型"""
    task_id: str
    status: str  # pending, processing, completed, failed
    progress: int
    message: str
    result: Optional[Dict[str, Any]] = None
    created_at: str
    updated_at: str
    error: Optional[str] = None

def update_task_status(task_id: str, status: str, progress: int, message: str, result: Dict[str, Any] = None, error: str = None):
    """更新任务状态"""
    if task_id in task_status:
        task_status[task_id].update({
            "status": status,
            "progress": progress,
            "message": message,
            "updated_at": datetime.now().isoformat(),
            "result": result,
            "error": error
        })

async def process_bid_proposal_generation_async(task_id: str, tender_document_json: Dict[str, Any], 
                                        model_name: str = None, batch_size: int = None):
    """异步处理投标书生成任务（不含附件）"""
    try:
        # 更新任务状态为处理中
        update_task_status(task_id, "processing", 20, "开始生成投标书...")
        
        # 不处理附件，直接生成投标书
        attachment_info = None
        
        # 初始化生成器
        generator = BidProposalGenerator(model_name=model_name)
        
        update_task_status(task_id, "processing", 30, "正在生成投标书...")
        
        # 生成投标书（包含附件信息）
        result = generator.generate_bid_proposal(
            tender_document_json=tender_document_json,
            model_name=model_name,
            batch_size=batch_size,
            attachment_info=attachment_info
        )
        
        if result["success"]:
            # 保存为多种格式的文档
            try:
                # 确保download目录及子目录存在
                base_download_dir = r"d:\Pyhton-learn\new\study\zhaobiaoshu\download"
                word_dir = os.path.join(base_download_dir, "word")
                markdown_dir = os.path.join(base_download_dir, "markdown")
                
                os.makedirs(word_dir, exist_ok=True)
                os.makedirs(markdown_dir, exist_ok=True)
                
                # 生成时间戳和基础文件名
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_filename = f"bid_proposal_{timestamp}"
                
                bid_content = result["bid_proposal"]
                word_filepath = None
                markdown_filepath = None
                
                # 1. 保存为Word文档
                try:
                    doc = Document()
                    doc.add_heading('投标书', 0)
                    
                    if isinstance(bid_content, str):
                        # 按段落分割内容
                        paragraphs = bid_content.split('\n\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                    elif isinstance(bid_content, dict):
                        # 如果是字典格式，按章节添加
                        for section_title, section_content in bid_content.items():
                            doc.add_heading(section_title, level=1)
                            if isinstance(section_content, str):
                                doc.add_paragraph(section_content)
                            elif isinstance(section_content, list):
                                for item in section_content:
                                    doc.add_paragraph(str(item))
                    
                    word_filename = f"{base_filename}.docx"
                    word_filepath = os.path.join(word_dir, word_filename)
                    doc.save(word_filepath)
                    logger.info(f"投标书已保存为Word文档: {word_filepath}")
                    
                except Exception as e:
                    logger.error(f"保存Word文档失败: {str(e)}")
                
                # 2. 保存为Markdown文档
                try:
                    markdown_filename = f"{base_filename}.md"
                    markdown_filepath = os.path.join(markdown_dir, markdown_filename)
                    
                    # 将内容转换为Markdown格式
                    markdown_content = _convert_to_markdown(bid_content)
                    
                    with open(markdown_filepath, 'w', encoding='utf-8') as f:
                        f.write(markdown_content)
                    
                    logger.info(f"投标书已保存为Markdown文档: {markdown_filepath}")
                    
                except Exception as e:
                    logger.error(f"保存Markdown文档失败: {str(e)}")
                
            except Exception as e:
                logger.error(f"保存文档失败: {str(e)}")
                word_filepath = None
                markdown_filepath = None
            
            # 转换结果格式以匹配前端期望
            formatted_result = {
                "bid_content": result["bid_proposal"],
                "sections_generated": result["statistics"].get("total_sections", 0),
                "total_pages": max(1, len(str(result["bid_proposal"])) // 2000),  # 估算页数
                "summary": f"投标书生成完成，共{result['statistics'].get('total_sections', 0)}个章节，耗时{result['statistics'].get('processing_duration', 0):.1f}秒",
                "docx_content": result["bid_proposal"],  # 直接使用文本内容，前端会处理
                "word_filename": os.path.basename(word_filepath) if word_filepath else None,  # 使用文件名替代完整路径
                "markdown_filename": os.path.basename(markdown_filepath) if markdown_filepath else None,  # 使用文件名替代完整路径
                "statistics": result["statistics"],
                "section_plan": result.get("section_plan", {})
            }
            update_task_status(task_id, "completed", 100, "投标书生成完成", formatted_result)
        else:
            error_msg = "未知错误"
            if isinstance(result, dict):
                error_msg = result.get("error", "未知错误")
            elif isinstance(result, str):
                error_msg = result
            update_task_status(task_id, "failed", 0, "投标书生成失败", None, error_msg)
            
    except Exception as e:
        logger.error(f"处理带附件的投标书生成任务时出错: {str(e)}")
        update_task_status(task_id, "failed", 0, "处理失败", None, str(e))

async def process_bid_proposal_with_attachments_async(task_id: str, tender_document_json: Dict[str, Any],
                                                    attachments: List[UploadFile],
                                                    model_name: str = None, 
                                                    batch_size: int = None,
                                                    generate_outline_only: bool = False):
    """异步处理带附件的投标书生成"""
    try:
        update_task_status(task_id, "processing", 5, "开始处理附件...")
        
        # ✅ 处理附件
        processed_attachments = None
        if attachments and any(att.filename for att in attachments):
            try:
                from .enhanced_attachment_processor import EnhancedAttachmentProcessor
                attachment_processor = EnhancedAttachmentProcessor()
                processed_attachments = await attachment_processor.process_files(attachments)
                update_task_status(task_id, "processing", 20, "附件处理完成，开始生成投标书...")
                logger.info(f"成功处理 {len(processed_attachments.get('files', []))} 个附件")
            except Exception as e:
                logger.error(f"处理附件失败: {str(e)}")
                update_task_status(task_id, "processing", 15, f"附件处理失败: {str(e)}，继续生成投标书...")
        else:
            update_task_status(task_id, "processing", 10, "无附件，开始生成投标书...")
        
        # 初始化生成器
        generator = BidProposalGenerator(model_name=model_name)
        
        update_task_status(task_id, "processing", 30, "正在生成投标书...")
        
        # ✅ 传入附件信息
        result = generator.generate_bid_proposal(
            tender_document_json=tender_document_json,
            model_name=model_name,
            batch_size=batch_size,
            attachment_info=processed_attachments  # ← 重要：传入附件信息
        )
        
        if result["success"]:
            # 保存为多种格式的文档
            try:
                # 确保download目录及子目录存在
                base_download_dir = r"d:\Pyhton-learn\new\study\zhaobiaoshu\download"
                word_dir = os.path.join(base_download_dir, "word")
                markdown_dir = os.path.join(base_download_dir, "markdown")
                
                os.makedirs(word_dir, exist_ok=True)
                os.makedirs(markdown_dir, exist_ok=True)
                
                # 生成时间戳和基础文件名
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_filename = f"bid_proposal_{timestamp}"
                
                bid_content = result["bid_proposal"]
                word_filepath = None
                markdown_filepath = None
                
                # 1. 保存为Word文档
                try:
                    doc = Document()
                    doc.add_heading('投标书', 0)
                    
                    if isinstance(bid_content, str):
                        # 按段落分割内容
                        paragraphs = bid_content.split('\n\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                    elif isinstance(bid_content, dict):
                        # 如果是字典格式，按章节添加
                        for section_title, section_content in bid_content.items():
                            doc.add_heading(section_title, level=1)
                            if isinstance(section_content, str):
                                doc.add_paragraph(section_content)
                            elif isinstance(section_content, list):
                                for item in section_content:
                                    doc.add_paragraph(str(item))
                    
                    word_filename = f"{base_filename}.docx"
                    word_filepath = os.path.join(word_dir, word_filename)
                    doc.save(word_filepath)
                    logger.info(f"投标书已保存为Word文档: {word_filepath}")
                    
                except Exception as e:
                    logger.error(f"保存Word文档失败: {str(e)}")
                
                # 2. 保存为Markdown文档
                try:
                    markdown_filename = f"{base_filename}.md"
                    markdown_filepath = os.path.join(markdown_dir, markdown_filename)
                    
                    # 将内容转换为Markdown格式
                    markdown_content = _convert_to_markdown(bid_content)
                    
                    with open(markdown_filepath, 'w', encoding='utf-8') as f:
                        f.write(markdown_content)
                    
                    logger.info(f"投标书已保存为Markdown文档: {markdown_filepath}")
                    
                except Exception as e:
                    logger.error(f"保存Markdown文档失败: {str(e)}")
                
            except Exception as e:
                logger.error(f"保存文档失败: {str(e)}")
                word_filepath = None
                markdown_filepath = None
            
            # 转换结果格式以匹配前端期望
            formatted_result = {
                "bid_content": result["bid_proposal"],
                "sections_generated": result["statistics"].get("total_sections", 0),
                "total_pages": max(1, len(str(result["bid_proposal"])) // 2000),  # 估算页数
                "summary": f"投标书生成完成，共{result['statistics'].get('total_sections', 0)}个章节，耗时{result['statistics'].get('processing_duration', 0):.1f}秒",
                "docx_content": result["bid_proposal"],  # 直接使用文本内容，前端会处理
                "word_filename": os.path.basename(word_filepath) if word_filepath else None,  # 使用文件名替代完整路径
                "markdown_filename": os.path.basename(markdown_filepath) if markdown_filepath else None,  # 使用文件名替代完整路径
                "statistics": result["statistics"],
                "section_plan": result.get("section_plan", {})
            }
            update_task_status(task_id, "completed", 100, "投标书生成完成", formatted_result)
        else:
            error_msg = "未知错误"
            if isinstance(result, dict):
                error_msg = result.get("error", "未知错误")
            elif isinstance(result, str):
                error_msg = result
            update_task_status(task_id, "failed", 0, "投标书生成失败", None, error_msg)
            
    except Exception as e:
        logger.error(f"处理投标书生成任务时出错: {str(e)}")
        update_task_status(task_id, "failed", 0, "处理失败", None, str(e))

def _convert_to_markdown(bid_content):
    """将投标书内容转换为Markdown格式
    
    改进的markdown格式化功能：
    1. 规范化标题层级
    2. 避免重复的一级标题
    3. 统一格式化处理
    4. 清理多余空行
    """
    try:
        import re
        from datetime import datetime
        
        markdown_content = "# 投标书\n\n"
        
        if isinstance(bid_content, str):
            # 字符串格式处理
            markdown_content += _process_string_content(bid_content)
                        
        elif isinstance(bid_content, dict):
            # 字典格式处理
            markdown_content += _process_dict_content(bid_content)
        else:
            # 其他格式，转换为字符串处理
            markdown_content += str(bid_content).strip() + "\n\n"
            
        # 清理格式
        markdown_content = _clean_markdown_format(markdown_content)
        
        # 添加生成时间信息
        markdown_content += "\n---\n\n"
        markdown_content += f"*本投标书由系统自动生成，生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}*\n"
        
        return markdown_content
        
    except Exception as e:
        logger.error(f"转换Markdown格式失败: {str(e)}")
        return f"# 投标书\n\n转换失败：{str(e)}\n\n{str(bid_content)}"


def _process_string_content(content):
    """处理字符串格式的投标书内容，优化标题处理"""
    import re
    
    result = ""
    paragraphs = content.split('\n\n')
    seen_titles = set()
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            continue
            
        paragraph = paragraph.strip()
        
        # 跳过重复的"投标书"标题和其他重复标题
        if paragraph in ["# 投标书", "投标书"] or paragraph.endswith(" 投标书"):
            continue
        
        # 处理已有的markdown标题
        if paragraph.startswith('#'):
            # 规范化标题格式
            title_level = len(paragraph) - len(paragraph.lstrip('#'))
            title_text = paragraph.lstrip('#').strip()
            
            if title_text:
                # 标准化标题
                normalized_title = _normalize_title(title_text)
                
                # 跳过重复标题
                if _is_duplicate_title(normalized_title, seen_titles):
                    continue
                
                # 验证和调整标题层级
                validated_level = _validate_title_level(normalized_title, title_level)
                # 确保不超过4级标题，一级标题改为二级
                if validated_level == 1:
                    validated_level = 2
                validated_level = min(validated_level, 4)
                
                seen_titles.add(normalized_title)
                result += f"{'#' * validated_level} {normalized_title}\n\n"
        
        # 检测可能的章节标题
        elif _is_chapter_title(paragraph):
            title = _clean_title_text(paragraph)
            if title:
                normalized_title = _normalize_title(title)
                
                # 跳过重复标题
                if not _is_duplicate_title(normalized_title, seen_titles):
                    seen_titles.add(normalized_title)
                    result += f"## {normalized_title}\n\n"
        
        # 检测可能的子标题
        elif _is_sub_title(paragraph):
            title = _clean_title_text(paragraph)
            if title:
                normalized_title = _normalize_title(title)
                
                # 跳过重复标题
                if not _is_duplicate_title(normalized_title, seen_titles):
                    seen_titles.add(normalized_title)
                    result += f"### {normalized_title}\n\n"
        
        else:
            # 普通段落
            result += paragraph + "\n\n"
    
    return result


def _process_dict_content(content_dict):
    """处理字典格式的投标书内容"""
    result = ""
    
    # 预定义的章节顺序
    section_order = [
        "基本信息", "项目理解", "第一章 投标人基本情况", "第二章 资格条件响应",
        "第三章 技术方案", "第四章 项目管理方案", "第五章 商务方案", 
        "第六章 售后服务方案", "第七章 其他"
    ]
    
    # 按预定义顺序处理章节
    processed_sections = set()
    
    for section_key in section_order:
        if section_key in content_dict:
            result += _format_section(section_key, content_dict[section_key])
            processed_sections.add(section_key)
    
    # 处理剩余章节
    for section_title, section_content in content_dict.items():
        if section_title not in processed_sections:
            result += _format_section(section_title, section_content)
    
    return result


def _format_section(title, content):
    """格式化单个章节"""
    result = f"## {title}\n\n"
    
    if isinstance(content, str):
        # 处理字符串内容中的子标题
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if _is_sub_title(line):
                clean_title = _clean_title_text(line)
                if clean_title:
                    result += f"### {clean_title}\n\n"
            else:
                result += line + "\n"
        result += "\n"
        
    elif isinstance(content, list):
        for item in content:
            if isinstance(item, str) and item.strip():
                result += f"- {item.strip()}\n"
        result += "\n"
    else:
        result += str(content).strip() + "\n\n"
    
    return result


def _is_chapter_title(text):
    """判断是否为章节标题，更精确的匹配"""
    import re
    
    text = text.strip()
    
    # 黑名单关键词 - 这些不应该被识别为章节标题
    blacklist_keywords = [
        '详见', '如下', '包括', '具体', '说明', '要求', '标准', '规范',
        '附件', '附录', '备注', '注意', '提醒', '温馨提示'
    ]
    
    # 检查黑名单
    for keyword in blacklist_keywords:
        if keyword in text:
            return False
    
    # 长度检查 - 太短或太长的文本不太可能是标题
    if len(text) < 2 or len(text) > 50:
        return False
    
    # 章节标题模式（更严格）
    chapter_patterns = [
        r'^第[一二三四五六七八九十\d]+章[\s：:].{2,30}$',  # 第X章 标题（限制长度）
        r'^[一二三四五六七八九十]、.{2,30}$',  # 一、标题
        r'^\d+[、.]\s*.{2,30}$',  # 1. 标题 或 1、标题
        r'^\d+\.\d+[\s：:].{2,30}$',  # 1.1 标题
        r'^.{2,20}[：:]$',  # 以冒号结尾的标题（限制长度）
    ]
    
    for pattern in chapter_patterns:
        if re.match(pattern, text):
            return True
    
    # 特定关键词（更精确匹配）
    title_keywords = [
        '基本信息', '项目理解', '技术方案', '商务方案', '服务方案', 
        '资格条件', '项目管理', '质量保证', '进度安排', '人员配置',
        '设备配置', '安全措施', '环保措施', '售后服务'
    ]
    
    # 关键词匹配需要更严格的条件
    for keyword in title_keywords:
        if keyword in text:
            # 确保关键词不是在句子中间
            if text.startswith(keyword) or text.endswith(keyword):
                return True
            # 或者关键词前后有标点符号
            if re.search(f'[^\\w]{keyword}[^\\w]', text):
                return True
    
    return False


def _is_sub_title(text):
    """判断是否为子标题"""
    import re
    
    # 子标题模式
    sub_patterns = [
        r'^\d+\.\d+\.\d+\s+.+',  # 1.1.1 标题
        r'^\([一二三四五六七八九十\d]+\)\s*.+',  # (1) 标题
        r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*.+',  # ① 标题
    ]
    
    for pattern in sub_patterns:
        if re.match(pattern, text.strip()):
            return True
    
    return False


def _clean_title_text(text):
    """清理标题文本"""
    import re
    
    # 移除标题标记
    text = re.sub(r'^#+\s*', '', text)  # 移除markdown标记
    text = re.sub(r'^第[一二三四五六七八九十\d]+章\s*', '', text)  # 移除章节标记
    text = re.sub(r'^[一二三四五六七八九十\d]+[、.]\s*', '', text)  # 移除序号
    text = re.sub(r'^\d+\.\d+\.?\s*', '', text)  # 移除数字序号
    text = re.sub(r'^\([一二三四五六七八九十\d]+\)\s*', '', text)  # 移除括号序号
    text = re.sub(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*', '', text)  # 移除圆圈序号
    text = text.rstrip('：:').strip()  # 移除结尾冒号
    
    return text.strip()


def _normalize_title(title):
    """标准化标题格式"""
    import re
    
    # 移除多余空白
    title = re.sub(r'\s+', ' ', title.strip())
    
    # 移除重复的标点符号
    title = re.sub(r'[：:]{2,}', '：', title)
    title = re.sub(r'[。.]{2,}', '。', title)
    
    # 统一标点符号
    title = title.replace(':', '：')
    
    return title


def _is_duplicate_title(title, seen_titles):
    """检查是否为重复标题"""
    normalized = _normalize_title(title)
    
    # 检查完全相同的标题
    if normalized in seen_titles:
        return True
    
    # 检查相似标题（去除标点后比较）
    import re
    clean_title = re.sub(r'[^\w\s]', '', normalized)
    for seen in seen_titles:
        clean_seen = re.sub(r'[^\w\s]', '', seen)
        if clean_title == clean_seen and clean_title:
            return True
    
    return False


def _validate_title_level(title, level):
    """验证标题层级是否合理"""
    import re
    
    # 一级标题关键词
    level1_keywords = ['投标书', '技术方案', '商务方案', '项目管理方案', '售后服务方案']
    
    # 二级标题关键词
    level2_keywords = ['基本信息', '项目理解', '资格条件', '技术要求', '服务承诺']
    
    # 检查是否包含章节标记
    if re.search(r'第[一二三四五六七八九十\d]+章', title):
        return min(level, 2)  # 章节标题最多为二级
    
    # 根据关键词调整级别
    for keyword in level1_keywords:
        if keyword in title:
            return 2  # 主要章节为二级标题
    
    for keyword in level2_keywords:
        if keyword in title:
            return 3  # 子章节为三级标题
    
    # 限制最大层级
    return min(level, 4)


def _clean_markdown_format(content):
    """清理markdown格式，消除标题重复和层级混乱"""
    import re
    
    lines = content.split('\n')
    cleaned_lines = []
    seen_titles = set()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            # 避免连续空行
            if cleaned_lines and cleaned_lines[-1] != '':
                cleaned_lines.append('')
            i += 1
            continue
        
        # 处理markdown标题
        title_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if title_match:
            level = len(title_match.group(1))
            title_text = title_match.group(2).strip()
            
            # 标准化标题
            normalized_title = _normalize_title(title_text)
            
            # 跳过重复标题
            if _is_duplicate_title(normalized_title, seen_titles):
                i += 1
                continue
            
            # 验证和调整标题层级
            validated_level = _validate_title_level(normalized_title, level)
            
            # 添加标题
            if normalized_title:
                seen_titles.add(normalized_title)
                # 确保标题前有空行（除非是第一行）
                if cleaned_lines and cleaned_lines[-1] != '':
                    cleaned_lines.append('')
                cleaned_lines.append(f"{'#' * validated_level} {normalized_title}")
                cleaned_lines.append('')  # 标题后空行
        else:
            # 普通内容行
            cleaned_lines.append(line)
        
        i += 1
    
    # 清理结尾多余空行
    while cleaned_lines and cleaned_lines[-1] == '':
        cleaned_lines.pop()
    
    return '\n'.join(cleaned_lines)

@router.post("/generate_from_json", response_model=BidProposalGenerationResponse, 
            summary="从招标文件JSON生成投标书", description="上传经过filter.py处理的招标文件JSON数据，生成完整投标书")
async def generate_bid_proposal_from_json(
    background_tasks: BackgroundTasks,
    request: BidProposalGenerationRequest
):
    """从招标文件JSON数据生成投标书"""
    try:
        logger.info("收到投标书生成请求")
        
        # 验证JSON数据
        if not request.tender_document_json:
            raise HTTPException(status_code=400, detail="招标文件JSON数据不能为空")
        
        if "content" not in request.tender_document_json:
            raise HTTPException(status_code=400, detail="招标文件JSON数据格式错误，缺少content字段")
        
        # 如果只生成大纲
        if request.generate_outline_only:
            generator = BidProposalGenerator(model_name=request.model_name)
            outline_result = generator.generate_bid_outline(request.tender_document_json)
            
            return BidProposalGenerationResponse(
                success=outline_result["success"],
                message="投标书大纲生成完成" if outline_result["success"] else "大纲生成失败",
                error=outline_result.get("error", ""),
                data=outline_result
            )
        
        # 生成任务ID
        task_id = str(uuid.uuid4())
        
        # 初始化任务状态
        task_status[task_id] = {
            "task_id": task_id,
            "status": "pending",
            "progress": 0,
            "message": "任务已创建，等待处理...",
            "result": None,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "error": None
        }
        
        # 添加后台任务
        background_tasks.add_task(
            process_bid_proposal_generation_async,
            task_id,
            request.tender_document_json,
            request.model_name,
            request.batch_size
        )
        
        return BidProposalGenerationResponse(
            success=True,
            task_id=task_id,
            message="任务已创建，正在处理中...",
            data={"status_url": f"/api/gender_book/status/{task_id}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建投标书生成任务失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")

@router.post("/upload_json", response_model=BidProposalGenerationResponse,
            summary="上传招标文件JSON生成投标书", description="上传招标文件JSON并生成投标书")
async def upload_json_and_generate(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(..., description="经过filter.py处理的招标文件JSON文件"),
    model_name: Optional[str] = Form(None, description="指定使用的模型名称"),
    batch_size: Optional[int] = Form(None, description="批处理大小"),
    generate_outline_only: Optional[bool] = Form(False, description="是否只生成大纲")
):
    """上传招标文件JSON并生成投标书"""
    try:
        # 验证文件格式
        if not file.filename or not file.filename.lower().endswith('.json'):
            raise HTTPException(status_code=400, detail="请上传JSON格式的文件")
        
        # 读取文件内容
        content = await file.read()
        
        try:
            json_data = json.loads(content.decode('utf-8'))
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"JSON文件格式错误: {str(e)}")
        except UnicodeDecodeError:
            raise HTTPException(status_code=400, detail="文件编码错误，请确保文件为UTF-8编码")
        
        # 创建请求对象
        request = BidProposalGenerationRequest(
            tender_document_json=json_data,
            model_name=model_name,
            batch_size=batch_size,
            generate_outline_only=generate_outline_only
        )
        
        # 调用生成函数
        return await generate_bid_proposal_from_json(background_tasks, request)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"处理文件上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"处理文件失败: {str(e)}")

@router.get("/status/{task_id}", response_model=TaskStatusResponse,
           summary="查询任务状态", description="查询投标书生成任务的状态")
async def get_task_status(task_id: str):
    """查询任务状态"""
    try:
        if task_id not in task_status:
            raise HTTPException(status_code=404, detail="任务不存在")
        
        status_info = task_status[task_id]
        
        return TaskStatusResponse(
            task_id=task_id,
            status=status_info["status"],
            progress=status_info["progress"],
            message=status_info["message"],
            result=status_info.get("result"),
            created_at=status_info["created_at"],
            updated_at=status_info["updated_at"],
            error=status_info.get("error")
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"查询任务状态失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")

@router.get("/tasks", summary="获取任务列表", description="获取所有任务的状态列表")
async def get_all_tasks(limit: int = 20, offset: int = 0):
    """获取任务列表"""
    try:
        # 按创建时间倒序排列
        sorted_tasks = sorted(
            task_status.values(),
            key=lambda x: x["created_at"],
            reverse=True
        )
        
        # 分页
        paginated_tasks = sorted_tasks[offset:offset + limit]
        
        return {
            "success": True,
            "total": len(task_status),
            "tasks": paginated_tasks,
            "has_more": offset + limit < len(task_status)
        }
        
    except Exception as e:
        logger.error(f"获取任务列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取任务列表失败: {str(e)}")

@router.delete("/tasks/{task_id}", summary="删除任务", description="删除指定的任务记录")
async def delete_task(task_id: str):
    """删除任务"""
    try:
        if task_id not in task_status:
            raise HTTPException(status_code=404, detail="任务不存在")
        
        del task_status[task_id]
        
        return {
            "success": True,
            "message": "任务删除成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除任务失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除任务失败: {str(e)}")

@router.post("/analyze_json", summary="分析JSON内容", description="分析JSON内容并返回章节计划，不生成具体内容")
async def analyze_json_content(request: Dict[str, Any]):
    """分析JSON内容"""
    try:
        logger.info("收到JSON内容分析请求")
        
        # 验证输入
        if not request or "json_data" not in request:
            raise HTTPException(status_code=400, detail="请求数据格式错误，缺少json_data字段")
        
        # 获取JSON数据
        json_data = request["json_data"]
        
        # 初始化章节管理器
        section_manager = SectionManager()
        
        # 分析内容
        content_analysis = section_manager.analyze_json_content(json_data)
        
        # 生成章节计划
        section_plan = section_manager.generate_section_plan(content_analysis)
        
        return {
            "success": True,
            "message": "JSON内容分析完成",
            "data": {
                "content_analysis": content_analysis,
                "section_plan": section_plan
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"上传JSON文件并生成投标书失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")

@router.post("/generate_from_json_with_attachments", response_model=BidProposalGenerationResponse,
            summary="从招标文件JSON生成投标书（支持附件）", 
            description="上传招标文件JSON和附件，生成包含附件说明的完整投标书")
async def generate_bid_proposal_with_attachments(
    background_tasks: BackgroundTasks,
    tender_document_json: str = Form(..., description="经过filter.py处理的招标文件JSON数据"),
    model_name: Optional[str] = Form(None, description="指定使用的模型名称"),
    batch_size: Optional[int] = Form(None, description="批处理大小"),
    generate_outline_only: Optional[bool] = Form(False, description="是否只生成大纲"),
    attachments: List[UploadFile] = File(None, description="附件文件列表（支持图片和PDF）")
):
    """从招标文件JSON数据和附件生成投标书"""
    try:
        logger.info("收到带附件的投标书生成请求")
        
        # 解析JSON数据
        try:
            json_data = json.loads(tender_document_json)
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"JSON数据格式错误: {str(e)}")
        
        # 验证JSON数据
        if not json_data:
            raise HTTPException(status_code=400, detail="招标文件JSON数据不能为空")
        
        if "content" not in json_data:
            raise HTTPException(status_code=400, detail="招标文件JSON数据格式错误，缺少content字段")
        
        # 生成任务ID
        task_id = str(uuid.uuid4())
        
        # 初始化任务状态
        task_status[task_id] = {
            "task_id": task_id,
            "status": "pending",
            "progress": 0,
            "message": "任务已创建，等待处理...",
            "result": None,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "error": None
        }
        
        # 添加后台任务（包含附件处理）
        background_tasks.add_task(
            process_bid_proposal_with_attachments_async,
            task_id,
            json_data,
            attachments,
            model_name,
            batch_size,
            generate_outline_only
        )
        
        return BidProposalGenerationResponse(
            success=True,
            task_id=task_id,
            message="投标书生成任务已创建（包含附件处理）",
            data={"status_url": f"/api/gender_book/status/{task_id}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"创建带附件的投标书生成任务失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")

@router.get("/health", summary="健康检查", description="检查gender_book模块的健康状态")
async def health_check():
    """健康检查"""
    try:
        # 测试章节管理器
        section_manager = SectionManager()
        
        # 测试数据
        test_data = {
            "content": [
                {"text": "测试项目", "type": "Title", "section": "项目概述"},
                {"text": "这是一个测试项目", "type": "Content", "section": "项目概述"}
            ]
        }
        
        # 执行分析测试
        content_analysis = section_manager.analyze_json_content(test_data)
        section_plan = section_manager.generate_section_plan(content_analysis)
        
        return {
            "status": "healthy",
            "module": "gender_book",
            "message": "投标书分章节生成模块运行正常",
            "features": {
                "section_analysis": "正常",
                "section_planning": "正常",
                "task_management": "正常"
            },
            "active_tasks": len(task_status),
            "standard_sections": len(section_manager.standard_sections)
        }
        
    except Exception as e:
        logger.error(f"健康检查失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"模块异常: {str(e)}")

@router.get("/sections", summary="获取标准章节模板", description="获取预定义的标准投标书章节模板")
async def get_standard_sections():
    """获取标准章节模板"""
    try:
        section_manager = SectionManager()
        
        sections_info = []
        for section_id, section in section_manager.standard_sections.items():
            sections_info.append({
                "id": section.id,
                "title": section.title,
                "description": section.description,
                "keywords": section.content_keywords,
                "priority": section.priority,
                "estimated_length": section.estimated_length
            })
        
        # 按优先级排序
        sections_info.sort(key=lambda x: x["priority"])
        
        return {
            "success": True,
            "total_sections": len(sections_info),
            "sections": sections_info
        }
        
    except Exception as e:
        logger.error(f"获取标准章节模板失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")

@router.get("/download/{file_type}/{filename}", summary="下载生成的文件", description="下载Word或Markdown格式的投标书文件")
async def download_file(file_type: str, filename: str):
    """下载生成的文件"""
    try:
        # 验证文件类型
        if file_type not in ["word", "markdown"]:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        
        # 构建文件路径
        base_download_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "download")
        file_dir = os.path.join(base_download_dir, file_type)
        file_path = os.path.join(file_dir, filename)
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 检查文件是否在允许的目录内（安全检查）
        if not os.path.abspath(file_path).startswith(os.path.abspath(file_dir)):
            raise HTTPException(status_code=403, detail="访问被拒绝")
        
        # 设置媒体类型
        media_type = "application/octet-stream"
        if file_type == "word":
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_type == "markdown":
            media_type = "text/markdown"
        
        # 返回文件
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type=media_type
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载文件失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"下载失败: {str(e)}")